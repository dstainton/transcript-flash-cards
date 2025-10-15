"""
Document processing module for extracting text from various file types.
Supports PDF, Word documents (.docx), and plain text files.
"""

import os
from typing import Dict, List, Optional
import PyPDF2
import pdfplumber
from docx import Document
import magic


class DocumentProcessor:
    """Handles text extraction from various document types"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
    
    def __init__(self):
        self.magic = magic.Magic(mime=True)
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if file type is supported"""
        ext = os.path.splitext(filename)[1].lower()
        return ext in self.SUPPORTED_EXTENSIONS
    
    def validate_file(self, file_path: str) -> tuple[bool, Optional[str]]:
        """
        Validate file size and type.
        Returns (is_valid, error_message)
        """
        # Check if file exists
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > self.MAX_FILE_SIZE:
            return False, f"File too large (max {self.MAX_FILE_SIZE // (1024*1024)}MB)"
        
        if file_size == 0:
            return False, "File is empty"
        
        # Check extension
        filename = os.path.basename(file_path)
        if not self.is_supported_file(filename):
            return False, f"Unsupported file type. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
        
        return True, None
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pdfplumber (better layout support)"""
        text = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
        except Exception as e:
            # Fallback to PyPDF2 if pdfplumber fails
            print(f"pdfplumber failed, trying PyPDF2: {e}")
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text.append(page_text)
            except Exception as e2:
                raise Exception(f"Failed to extract PDF text: {e2}")
        
        return '\n\n'.join(text)
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            doc = Document(file_path)
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text.append(' | '.join(row_text))
            
            return '\n\n'.join(text)
        except Exception as e:
            raise Exception(f"Failed to extract Word document text: {e}")
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Unable to decode text file with supported encodings")
        except Exception as e:
            raise Exception(f"Failed to extract text file content: {e}")
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from document based on file type.
        Returns the extracted text or raises an exception.
        """
        # Validate file first
        is_valid, error_msg = self.validate_file(file_path)
        if not is_valid:
            raise ValueError(error_msg)
        
        # Get file extension
        ext = os.path.splitext(file_path)[1].lower()
        
        # Extract text based on file type
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def process_multiple_documents(self, file_paths: List[str]) -> Dict[str, str]:
        """
        Process multiple documents and return a dict of {filename: extracted_text}.
        Skips files that fail processing and logs errors.
        """
        results = {}
        errors = {}
        
        for file_path in file_paths:
            filename = os.path.basename(file_path)
            try:
                text = self.extract_text(file_path)
                results[filename] = text
                print(f"✓ Successfully processed: {filename}")
            except Exception as e:
                errors[filename] = str(e)
                print(f"✗ Failed to process {filename}: {e}")
        
        return results, errors

